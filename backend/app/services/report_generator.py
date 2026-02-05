from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime, timezone
from io import BytesIO
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


@dataclass
class GeneratedReport:
    filename: str
    mime_type: str
    content: bytes


def _safe_str(v: Any) -> str:
    if v is None:
        return ''
    if isinstance(v, (list, tuple)):
        return ', '.join([str(x) for x in v if x is not None])
    return str(v)


def build_seat_report_docx(
    project: Dict[str, Any],
    documents: List[Dict[str, Any]],
    *,
    extracted: Optional[Dict[str, Any]] = None,
) -> GeneratedReport:
    """Create a DOCX report using the user's provided template skeleton.

    The skeleton is based on the "Seat of Arbitration Selection Template".
    We keep the structure stable so later you can replace placeholders with your
    real decision engine outputs.
    """

    extracted = extracted or {}
    intake = project.get('intake') or {}

    doc = Document()

    # Global style tweaks (kept light to avoid fighting Word defaults)
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    title = project.get('title') or 'Seat of Arbitration Selection'
    doc.add_heading(title, level=0)

    # Generated timestamp
    ts = datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')
    p = doc.add_paragraph(f'Generated: {ts}')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Project details table
    doc.add_paragraph('')
    doc.add_heading('Project details', level=2)
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Light Shading Accent 1'

    def add_row(k: str, v: Any):
        row = table.add_row().cells
        row[0].text = k
        row[1].text = _safe_str(v) if _safe_str(v) else '—'

    add_row('Project ID', project.get('id'))
    add_row('Status', project.get('status'))
    add_row('Current seat', intake.get('current_seat') or extracted.get('current_seat'))
    add_row('Proposed seat(s)', intake.get('proposed_seats') or extracted.get('proposed_seats'))
    add_row('Institution / rules', intake.get('institution_rules') or extracted.get('institution_rules'))
    add_row('Governing law (contract)', intake.get('governing_law') or extracted.get('governing_law'))
    add_row('Urgency (interim measures)', intake.get('urgency') or extracted.get('urgency'))
    add_row('Parties / assets location', intake.get('parties_assets_where') or extracted.get('parties_assets_where'))

    clause = intake.get('arbitration_agreement_text') or extracted.get('arbitration_agreement_text')
    if clause:
        add_row('Arbitration agreement (excerpt)', clause[:500] + ('…' if len(clause) > 500 else ''))

    doc.add_paragraph('')
    doc.add_heading('Sources provided', level=2)
    if documents:
        for d in documents:
            name = d.get('filename') or 'document'
            meta = []
            if d.get('mime_type'):
                meta.append(d['mime_type'])
            if d.get('byte_size'):
                meta.append(f"{int(d['byte_size'])} bytes")
            line = name + (f" ({', '.join(meta)})" if meta else '')
            doc.add_paragraph(line, style='List Bullet')
    else:
        doc.add_paragraph('No documents were uploaded at the time this report was generated.')

    doc.add_paragraph('')
    doc.add_paragraph(
        'Note: This document is an automatically generated working draft. It may contain placeholders when '
        'key information is missing. Review extracted details before relying on them.'
    )

    doc.add_paragraph('Disclaimer: This output is for decision-support and demonstration only and is not legal advice. It may be incomplete or incorrect. Always have a qualified practitioner review the underlying documents and applicable rules.')

    # Template skeleton sections
    doc.add_page_break()

    doc.add_heading('1. Grounds of the Dispute', level=1)
    doc.add_heading('1.1 Nature of the Dispute', level=2)
    doc.add_paragraph(extracted.get('nature_of_dispute') or intake.get('nature_of_dispute') or 'Not provided.')

    doc.add_heading('1.2 Parties', level=2)
    parties = extracted.get('parties') or intake.get('parties')
    if parties:
        if isinstance(parties, list):
            for x in parties:
                doc.add_paragraph(_safe_str(x), style='List Bullet')
        else:
            doc.add_paragraph(_safe_str(parties))
    else:
        doc.add_paragraph('Not provided.')

    doc.add_heading('1.3 Governing Law and Arbitration Agreement', level=2)
    gl = intake.get('governing_law') or extracted.get('governing_law') or extracted.get('governing_law')
    if gl:
        doc.add_paragraph(f'Governing law: {_safe_str(gl)}')
    else:
        doc.add_paragraph('Governing law: Not provided.')

    if clause:
        doc.add_paragraph('Arbitration agreement text (as provided):')
        doc.add_paragraph(_safe_str(clause))
    else:
        doc.add_paragraph('Arbitration agreement text: Not provided.')

    doc.add_heading('1.4 Procedural Sensitivities', level=2)
    proc = extracted.get('procedural_sensitivities') or intake.get('procedural_sensitivities')
    if proc:
        if isinstance(proc, list):
            for x in proc:
                doc.add_paragraph(_safe_str(x), style='List Bullet')
        else:
            doc.add_paragraph(_safe_str(proc))
    else:
        doc.add_paragraph('Not provided.')

    doc.add_heading('2. Compatible Jurisdictions', level=1)
    doc.add_heading('2.1 Selection Criteria', level=2)
    criteria = extracted.get('selection_criteria') or []
    if criteria:
        for x in criteria:
            doc.add_paragraph(_safe_str(x), style='List Bullet')
    else:
        for x in [
            'Enforceability and annulment risk considerations',
            'Court support vs interference',
            'Interim measures availability and urgency',
            'Compatibility with institution rules and clause wording',
            'Practicalities (language, logistics), where relevant',
        ]:
            doc.add_paragraph(x, style='List Bullet')

    doc.add_heading('2.2 Shortlisted Jurisdictions', level=2)
    shortlisted = extracted.get('shortlisted_jurisdictions')
    if shortlisted:
        if isinstance(shortlisted, list):
            for x in shortlisted:
                doc.add_paragraph(_safe_str(x), style='List Bullet')
        else:
            doc.add_paragraph(_safe_str(shortlisted))
    else:
        proposed = intake.get('proposed_seats') or extracted.get('proposed_seats')
        if proposed:
            for x in (proposed if isinstance(proposed, list) else [proposed]):
                doc.add_paragraph(_safe_str(x), style='List Bullet')
        else:
            doc.add_paragraph('No proposed seats were provided.')

    doc.add_heading('2.3 Jurisdiction-Specific Notes', level=2)
    notes = extracted.get('jurisdiction_notes')
    if notes and isinstance(notes, dict):
        for seat, v in notes.items():
            # Accept either a plain string or {pros:[], cons:[]}
            if isinstance(v, dict):
                pros = v.get('pros') or []
                cons = v.get('cons') or []
                if not isinstance(pros, list):
                    pros = [str(pros)]
                if not isinstance(cons, list):
                    cons = [str(cons)]
                txt = "Pros: " + "; ".join([_safe_str(x) for x in pros if x])
                if cons:
                    txt += "\nCons: " + "; ".join([_safe_str(x) for x in cons if x])
                doc.add_paragraph(f"{seat}: {txt.strip()}", style='List Bullet')
            else:
                doc.add_paragraph(f"{seat}: {_safe_str(v)}", style='List Bullet')
    else:
        doc.add_paragraph(
            'To be completed during the analysis phase. Add evidence-backed positives/negatives per seat here.'
        )

    doc.add_heading('3. Conclusion – Most Ideal Seat of Arbitration', level=1)
    doc.add_heading('3.1 Preferred Seat', level=2)
    preferred = extracted.get('preferred_seat')
    decision_flag = extracted.get('should_change_seat')
    if decision_flag:
        pretty_flag = {'yes': 'Yes', 'no': 'No', 'intervention_required': 'Intervention required'}.get(str(decision_flag), str(decision_flag))
        doc.add_paragraph(f'Seat change recommendation: {pretty_flag}')
    doc.add_paragraph('Recommended seat of arbitration:')
    doc.add_paragraph(_safe_str(preferred) if preferred else 'To be determined.')

    missing = extracted.get('missing_info')
    if decision_flag == 'intervention_required' and missing:
        doc.add_paragraph('Missing information / blockers:')
        if not isinstance(missing, list):
            missing = [str(missing)]
        for x in missing:
            doc.add_paragraph(_safe_str(x), style='List Bullet')

    doc.add_heading('3.2 Rationale', level=2)
    rationale = extracted.get('rationale')
    if rationale:
        if isinstance(rationale, list):
            for x in rationale:
                doc.add_paragraph(_safe_str(x), style='List Bullet')
        else:
            doc.add_paragraph(_safe_str(rationale))
    else:
        doc.add_paragraph('To be completed once the seat comparison matrix and evidence review are available.')


    # Optional: include a short citation list for transparency
    citations = extracted.get('citations')
    if citations and isinstance(citations, list):
        doc.add_paragraph('Evidence citations (excerpts):')
        for c in citations[:20]:
            if not isinstance(c, dict):
                continue
            src = str(c.get('source') or '')
            page = c.get('page')
            quote = _safe_str(c.get('quote'))
            if len(quote) > 180:
                quote = quote[:180] + '…'
            label = src
            if page:
                label += f" p.{page}"
            if quote:
                label += f": {quote}"
            doc.add_paragraph(label.strip(), style='List Bullet')

    doc.add_heading('3.3 Alternative or Backup Seat (if applicable)', level=2)
    alt = extracted.get('alternative_seat')
    doc.add_paragraph('Secondary recommended seat:')
    doc.add_paragraph(_safe_str(alt) if alt else 'To be determined.')
    doc.add_paragraph('Circumstances in which this alternative may be preferable:')
    doc.add_paragraph(
        _safe_str(extracted.get('alternative_circumstances'))
        if extracted.get('alternative_circumstances')
        else 'To be determined.'
    )

    bio = BytesIO()
    doc.save(bio)
    content = bio.getvalue()

    return GeneratedReport(
        filename=f"{(project.get('title') or 'seat-report').replace('/', '_')}.docx",
        mime_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        content=content,
    )
